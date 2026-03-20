"""
DOCX Generator for academic papers.

Converts paper sections to a formatted Microsoft Word document.
Supports both Chinese (GB/T7714) and English (APA/MLA/IEEE) formatting styles.
Also supports PDF conversion.
"""

from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path
import subprocess
import platform
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

import re
import asyncio
from skills.image_search import ImageSearchSkill, ChartGenerator
from skills.chart_generator import ChartGeneratorSkill


class DocxGenerator:
    """Generator for academic paper DOCX files."""

    # Format configurations for different citation styles
    FORMAT_CONFIGS = {
        'chinese': {
            'font_name': '宋体',
            'font_name_latin': 'Times New Roman',  # For English text in Chinese papers
            'font_name_bold': '黑体',  # Heiti for headings
            'title_size': 18,  # 小二 ~18pt
            # Chinese academic paper format standards
            'heading1_size': 14,  # 四号 ~14pt (一级标题)
            'heading2_size': 12,  # 小四 ~12pt (二级标题)
            'heading3_size': 12,  # 小四 ~12pt (三级标题)
            'body_size': 12,  # 小四 ~12pt
            'abstract_title_size': 15,  # 小三 ~15pt (摘要标题)
            'abstract_size': 12,  # 小四 ~12pt (摘要正文)
            'keyword_label_size': 14,  # 四号 ~14pt (关键词标签)
            'keyword_content_size': 12,  # 小四 ~12pt (关键词内容)
            'line_spacing': 1.5,  # 1.5倍行距
            'first_line_indent': Cm(0.74),  # ~2 Chinese characters (2字符)
            'margin_top': Cm(2.54),
            'margin_bottom': Cm(2.54),
            'margin_left': Cm(3.17),  # Chinese standard: 3.17cm left
            'margin_right': Cm(3.17),
        },
        'english': {
            'font_name': 'Times New Roman',
            'title_size': 18,
            'heading1_size': 16,
            'heading2_size': 14,
            'body_size': 12,
            'abstract_size': 11,
            'line_spacing': 1.5,  # 1.5 line spacing
            'first_line_indent': Inches(0.5),
            'margin_top': Inches(1),
            'margin_bottom': Inches(1),
            'margin_left': Inches(1.25),
            'margin_right': Inches(1.25),
        }
    }

    def __init__(self):
        """Initialize the DOCX generator with default styles."""
        self.format_config = None
        self.styles = {}

    def _setup_format(self, metadata: Dict[str, Any]):
        """Setup format configuration based on citation style or language."""
        citation_style = metadata.get('citation_style', 'apa').lower()
        language = metadata.get('language', 'english').lower()

        # Determine format: Chinese papers use GB/T7714 style
        if citation_style == 'gbt7714' or language == 'chinese':
            self.format_config = self.FORMAT_CONFIGS['chinese']
            self.is_chinese = True
        else:
            self.format_config = self.FORMAT_CONFIGS['english']
            self.is_chinese = False

        # Setup styles based on format
        if self.is_chinese:
            self.styles = {
                'title': {
                    'font_name': self.format_config['font_name'],
                    'font_size': self.format_config['title_size'],
                    'bold': True,
                    'alignment': WD_ALIGN_PARAGRAPH.CENTER,
                },
                'heading1': {
                    'font_name': self.format_config['font_name_bold'],  # 黑体
                    'font_size': self.format_config['heading1_size'],  # 四号
                    'bold': True,
                    'alignment': WD_ALIGN_PARAGRAPH.LEFT,
                },
                'heading2': {
                    'font_name': self.format_config['font_name_bold'],  # 黑体
                    'font_size': self.format_config['heading2_size'],  # 小四
                    'bold': True,
                    'alignment': WD_ALIGN_PARAGRAPH.LEFT,
                },
                'heading3': {
                    'font_name': self.format_config['font_name'],  # 宋体
                    'font_size': self.format_config['heading3_size'],  # 小四
                    'bold': False,
                    'alignment': WD_ALIGN_PARAGRAPH.LEFT,
                },
                'normal': {
                    'font_name': self.format_config['font_name'],
                    'font_size': self.format_config['body_size'],
                    'bold': False,
                    'alignment': WD_ALIGN_PARAGRAPH.LEFT,
                },
                'abstract_title': {
                    'font_name': self.format_config['font_name_bold'],  # 黑体
                    'font_size': self.format_config['abstract_title_size'],  # 小三
                    'bold': True,
                    'alignment': WD_ALIGN_PARAGRAPH.CENTER,
                },
                'abstract': {
                    'font_name': self.format_config['font_name'],  # 宋体
                    'font_size': self.format_config['abstract_size'],  # 小四
                    'bold': False,
                    'alignment': WD_ALIGN_PARAGRAPH.LEFT,
                    'italic': False,
                },
                'keyword_label': {
                    'font_name': self.format_config['font_name_bold'],  # 黑体
                    'font_size': self.format_config['keyword_label_size'],  # 四号
                    'bold': True,
                },
                'keyword_content': {
                    'font_name': self.format_config['font_name'],  # 宋体
                    'font_size': self.format_config['keyword_content_size'],  # 小四
                    'bold': False,
                },
            }
        else:
            self.styles = {
                'title': {
                    'font_name': self.format_config['font_name'],
                    'font_size': self.format_config['title_size'],
                    'bold': True,
                    'alignment': WD_ALIGN_PARAGRAPH.CENTER,
                },
                'heading1': {
                    'font_name': self.format_config['font_name'],
                    'font_size': self.format_config['heading1_size'],
                    'bold': True,
                    'alignment': WD_ALIGN_PARAGRAPH.LEFT,
                },
                'heading2': {
                    'font_name': self.format_config['font_name'],
                    'font_size': self.format_config['heading2_size'],
                    'bold': True,
                    'alignment': WD_ALIGN_PARAGRAPH.LEFT,
                },
                'heading3': {
                    'font_name': self.format_config['font_name'],
                    'font_size': self.format_config['body_size'],
                    'bold': True,
                    'alignment': WD_ALIGN_PARAGRAPH.LEFT,
                },
                'normal': {
                    'font_name': self.format_config['font_name'],
                    'font_size': self.format_config['body_size'],
                    'bold': False,
                    'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
                },
                'abstract_title': {
                    'font_name': self.format_config['font_name'],
                    'font_size': self.format_config['heading1_size'],
                    'bold': True,
                    'alignment': WD_ALIGN_PARAGRAPH.LEFT,
                },
                'abstract': {
                    'font_name': self.format_config['font_name'],
                    'font_size': self.format_config['abstract_size'],
                    'bold': False,
                    'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
                    'italic': True,
                },
            }

    def generate_paper(
        self,
        sections: Dict[str, str],
        metadata: Dict[str, Any],
        output_path: Optional[Path] = None,
        figures: Optional[List[Dict]] = None,
        paper_id: Optional[str] = None
    ) -> Path:
        """
        Generate a DOCX file from paper sections.

        Args:
            sections: Dictionary with paper sections
            metadata: Paper metadata (title, keywords, field, etc.)
            output_path: Optional output path, defaults to output/{paper_id}.docx
            figures: Optional list of figure definitions to include
            paper_id: Paper ID for organizing images

        Returns:
            Path to the generated DOCX file
        """
        self.paper_id = paper_id or "default"
        self._generated_figures = []  # Track which figures have been rendered
        # Setup format based on citation style/language
        self._setup_format(metadata)

        doc = Document()

        # Set up document margins (standard academic format)
        sections_doc = doc.sections[0]
        sections_doc.page_height = Inches(11.69)  # A4
        sections_doc.page_width = Inches(8.27)    # A4
        sections_doc.top_margin = self.format_config['margin_top']
        sections_doc.bottom_margin = self.format_config['margin_bottom']
        sections_doc.left_margin = self.format_config['margin_left']
        sections_doc.right_margin = self.format_config['margin_right']

        # Add title
        self._add_title(doc, sections.get('title', 'Untitled Paper'))

        # Add metadata
        self._add_metadata(doc, metadata)

        # Add horizontal line after metadata (only for English papers)
        if not self.is_chinese:
            doc.add_paragraph('_' * 50)

        # Add sections
        if self.is_chinese:
            # Chinese section headings
            section_order = [
                ('abstract', '摘要'),
                ('introduction', '1 引言'),
                ('literature_review', '2 文献综述'),
                ('methodology', '3 研究方法'),
                ('results', '4 研究结果'),
                ('discussion', '5 讨论'),
                ('conclusion', '6 结论'),
                ('references', '参考文献'),
            ]
        else:
            # English section headings
            section_order = [
                ('abstract', 'Abstract'),
                ('introduction', '1. Introduction'),
                ('literature_review', '2. Literature Review'),
                ('methodology', '3. Methodology'),
                ('results', '4. Results'),
                ('discussion', '5. Discussion'),
                ('conclusion', '6. Conclusion'),
                ('references', 'References'),
            ]

        for key, heading in section_order:
            content = sections.get(key)
            if content:
                self._add_section(doc, heading, content, key, figures)

        # Save document
        if output_path is None:
            output_dir = Path('output')
            output_dir.mkdir(exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = output_dir / f"paper_{timestamp}.docx"

        doc.save(output_path)
        return output_path

    def _add_title(self, doc: Document, title: str):
        """Add the paper title."""
        paragraph = doc.add_paragraph()
        paragraph.alignment = self.styles['title']['alignment']

        run = paragraph.add_run(title)
        run.bold = self.styles['title']['bold']
        run.font.name = self.styles['title']['font_name']
        run.font.size = Pt(self.styles['title']['font_size'])

        # For Chinese papers, also set East Asian font
        if self.is_chinese:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles['title']['font_name'])

        # Add spacing after title
        doc.add_paragraph()

    def _add_metadata(self, doc: Document, metadata: Dict[str, Any]):
        """Add paper metadata (keywords, field, citation style)."""
        language = metadata.get('language', 'english')

        # Keywords
        keywords = metadata.get('keywords', [])
        if keywords:
            if self.is_chinese:
                # 中文关键词格式：四号黑体标签，小四号宋体/ Times New Roman 内容，关键词之间用分号分隔
                para = doc.add_paragraph()
                # 关键词标签（四号黑体）
                label_run = para.add_run('关键词：')
                label_run.bold = True
                label_run.font.name = self.format_config['font_name_bold']  # 黑体
                label_run.font.size = Pt(self.format_config['keyword_label_size'])  # 四号 ~14pt
                label_run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_config['font_name_bold'])

                # 关键词内容（小四号宋体/ Times New Roman）
                keywords_text = '; '.join(keywords)
                self._add_mixed_text_run(para, keywords_text, self.format_config['keyword_content_size'],
                                          self.format_config['font_name'])
            else:
                label = 'Keywords:'
                para = doc.add_paragraph()
                run = para.add_run(f"{label} {', '.join(keywords)}")
                self._apply_font(run, 'normal', is_small=True)
                run.italic = True

        # Field - only for non-Chinese papers (中文论文通常不需要显示领域)
        if not self.is_chinese:
            field = metadata.get('field', '')
            if field:
                label = 'Field:'
                para = doc.add_paragraph()
                run = para.add_run(f"{label} {field}")
                self._apply_font(run, 'normal', is_small=True)

            # Citation Style
            citation_style = metadata.get('citation_style', 'apa')
            if citation_style:
                label = 'Citation Style:'
                para = doc.add_paragraph()
                run = para.add_run(f"{label} {citation_style.upper()}")
                self._apply_font(run, 'normal', is_small=True)

            # Language
            if language:
                label = 'Language:'
                lang_display = 'English'
                para = doc.add_paragraph()
                run = para.add_run(f"{label} {lang_display}")
                self._apply_font(run, 'normal', is_small=True)

    def _apply_font(self, run, style_key: str, is_small: bool = False, is_abstract: bool = False):
        """Apply font settings to a run, handling Chinese/English fonts."""
        style = self.styles.get(style_key, self.styles['normal'])

        if is_small:
            run.font.size = Pt(10)
        elif is_abstract:
            run.font.size = Pt(self.format_config['abstract_size'])
        else:
            run.font.size = Pt(style['font_size'])

        run.font.name = style['font_name']

        # For Chinese papers, set East Asian font
        if self.is_chinese:
            font_name = style['font_name']
            # Use Times New Roman as fallback for Latin characters in Chinese papers
            if style_key == 'normal' and 'font_name_latin' in self.format_config:
                run.font.name = self.format_config['font_name_latin']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        if 'italic' in style and style['italic'] and not self.is_chinese:
            run.italic = True
        if 'bold' in style:
            run.bold = style['bold']

    def _add_section(self, doc: Document, heading: str, content: str, section_key: str, figures: List[Dict] = None):
        """Add a section with heading and content."""
        # Special handling for abstract section in Chinese papers
        if self.is_chinese and section_key == 'abstract':
            self._add_chinese_abstract(doc, content)
            return

        # Special handling for references section in Chinese papers
        if self.is_chinese and section_key == 'references':
            self._add_chinese_references(doc, heading, content)
            return

        # Add heading
        heading_para = doc.add_paragraph()

        # Determine heading level for Chinese papers
        if self.is_chinese:
            heading_para.alignment = self.styles['heading1']['alignment']
            # Check heading level from heading text
            if heading.startswith(('1 ', '2 ', '3 ', '4 ', '5 ', '6 ', '7 ', '8 ', '9 ')):
                # Level 1 heading: "1 引言"
                style_key = 'heading1'
            elif heading.startswith('（') and heading[1].isdigit():
                # Level 2 heading: "（一）"
                style_key = 'heading2'
            elif heading[0].isdigit() and '. ' in heading:
                # Level 3 heading: "1. xxx"
                style_key = 'heading3'
            else:
                style_key = 'heading1'

            run = heading_para.add_run(heading)
            run.bold = self.styles[style_key]['bold']
            run.font.name = self.styles[style_key]['font_name']
            run.font.size = Pt(self.styles[style_key]['font_size'])
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles[style_key]['font_name'])

            # Set line spacing for heading
            heading_para.paragraph_format.line_spacing = self.format_config['line_spacing']
        else:
            heading_para.alignment = self.styles['heading1']['alignment']
            run = heading_para.add_run(heading)
            run.bold = self.styles['heading1']['bold']
            run.font.name = self.styles['heading1']['font_name']
            run.font.size = Pt(self.styles['heading1']['font_size'])

        # Add spacing after heading
        heading_para.paragraph_format.space_after = Pt(12)

        # Add content
        # Split content by paragraphs
        paragraphs = content.strip().split('\n\n')

        print(f"Processing section '{heading}' with {len(paragraphs)} paragraphs, figures: {len(figures) if figures else 0}")

        for para_text in paragraphs:
            # Check for figure markers
            figure_match = re.match(r'\[FIGURE:(\d+):([^:]+):([^\]]+)\]', para_text.strip())
            if figure_match:
                # Extract figure info
                fig_num = int(figure_match.group(1))
                fig_type = figure_match.group(2)
                fig_title = figure_match.group(3)

                # Find figure in figures list
                if figures:
                    fig_data = next((f for f in figures if f.get('number') == fig_num), None)
                    if fig_data:
                        self._add_figure(doc, fig_data)
                continue
            para_text = para_text.strip()
            if not para_text:
                continue

            # Check if it's a list item
            if para_text.startswith('- ') or para_text.startswith('* '):
                self._add_list_paragraph(doc, para_text, section_key)
            else:
                self._add_normal_paragraph(doc, para_text, section_key)

        # Add spacing after section
        doc.add_paragraph()

    def _add_chinese_abstract(self, doc: Document, content: str):
        """Add Chinese abstract with proper formatting."""
        # Add "摘要" title (小三号，黑体，加粗，居中)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run('摘要')
        run.bold = True
        run.font.name = self.format_config['font_name_bold']  # 黑体
        run.font.size = Pt(self.format_config['abstract_title_size'])  # 小三 ~15pt
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_config['font_name_bold'])

        # Add spacing after title
        title_para.paragraph_format.space_after = Pt(6)

        # Add abstract content (小四号，宋体/ Times New Roman，1.5倍行距，首行缩进2字符)
        content_para = doc.add_paragraph()
        content_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        content_para.paragraph_format.line_spacing = self.format_config['line_spacing']
        content_para.paragraph_format.first_line_indent = self.format_config['first_line_indent']

        # Use mixed font for abstract content
        self._add_mixed_text_run(content_para, content, self.format_config['abstract_size'],
                                  self.format_config['font_name'])

        # Add spacing after abstract
        doc.add_paragraph()

    def _add_chinese_references(self, doc: Document, heading: str, content: str):
        """Add Chinese references with proper formatting."""
        # Add "参考文献" title (四号黑体加粗，居中)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(heading)
        run.bold = True
        run.font.name = self.format_config['font_name_bold']  # 黑体
        run.font.size = Pt(self.format_config['heading1_size'])  # 四号 ~14pt
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.format_config['font_name_bold'])

        # Add spacing after title
        title_para.paragraph_format.space_after = Pt(6)

        # Add references content (小四号宋体/ Times New Roman，1.5倍行距)
        for ref in content.split('\n'):
            ref = ref.strip()
            if not ref:
                continue

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = self.format_config['line_spacing']
            # No first line indent for references

            # Use mixed font for references
            self._add_mixed_text_run(para, ref, self.format_config['body_size'],
                                      self.format_config['font_name'])

    def _add_normal_paragraph(self, doc: Document, text: str, section_key: str):
        """Add a normal paragraph with proper formatting for Chinese/English."""
        para = doc.add_paragraph()
        para.alignment = self.styles['normal']['alignment']

        # Set line spacing and indentation based on format
        if self.is_chinese:
            # Chinese: 1.5倍行距, 2-character first line indent
            para.paragraph_format.line_spacing = self.format_config['line_spacing']
            para.paragraph_format.first_line_indent = self.format_config['first_line_indent']
        else:
            # English: 1.5 line spacing, 0.5 inch first line indent
            para.paragraph_format.line_spacing = self.format_config['line_spacing']
            para.paragraph_format.first_line_indent = self.format_config['first_line_indent']

        # Use abstract style for abstract section
        is_abstract = section_key == 'abstract'
        style = self.styles['abstract'] if is_abstract else self.styles['normal']

        # For Chinese papers, split text to apply different fonts for Chinese and English/numbers
        if self.is_chinese:
            self._add_mixed_text_run(para, text, style['font_size'], style['font_name'])
        else:
            run = para.add_run(text)
            run.font.name = style['font_name']
            run.font.size = Pt(style['font_size'])
            if 'italic' in style and style['italic']:
                run.italic = style['italic']

    def _add_mixed_text_run(self, para, text: str, font_size: int, chinese_font: str):
        """
        Add text with mixed Chinese/English content, applying appropriate fonts.
        Chinese characters use chinese_font, English/numbers use Times New Roman.
        """
        import re

        # Pattern to match Chinese characters
        chinese_pattern = re.compile(r'[\u4e00-\u9fff]')

        current_run_text = ""
        is_current_chinese = None

        for char in text:
            is_char_chinese = bool(chinese_pattern.match(char))

            # If type changed (or first char), flush current run
            if is_current_chinese is not None and is_char_chinese != is_current_chinese:
                if current_run_text:
                    run = para.add_run(current_run_text)
                    run.font.size = Pt(font_size)
                    if is_current_chinese:
                        # Chinese font
                        run.font.name = chinese_font
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                    else:
                        # English/number font - Times New Roman
                        run.font.name = 'Times New Roman'
                current_run_text = char
                is_current_chinese = is_char_chinese
            else:
                current_run_text += char
                if is_current_chinese is None:
                    is_current_chinese = is_char_chinese

        # Flush remaining text
        if current_run_text:
            run = para.add_run(current_run_text)
            run.font.size = Pt(font_size)
            if is_current_chinese:
                run.font.name = chinese_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
            else:
                run.font.name = 'Times New Roman'

    def _add_list_paragraph(self, doc: Document, text: str, section_key: str):
        """Add a list item paragraph."""
        # Remove list marker
        clean_text = text.lstrip('- *').strip()

        para = doc.add_paragraph(style='List Bullet')

        if self.is_chinese:
            para.paragraph_format.left_indent = Cm(0.74)  # ~2 Chinese characters
            para.paragraph_format.first_line_indent = Cm(-0.37)  # ~1 Chinese character
            para.paragraph_format.line_spacing = self.format_config['line_spacing']  # 1.5倍行距
        else:
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            para.paragraph_format.line_spacing = self.format_config['line_spacing']

        if self.is_chinese:
            # Use mixed font for list items
            self._add_mixed_text_run(para, clean_text, self.styles['normal']['font_size'],
                                      self.styles['normal']['font_name'])
        else:
            run = para.add_run(clean_text)
            run.font.name = self.styles['normal']['font_name']
            run.font.size = Pt(self.styles['normal']['font_size'])

    def _add_figure(self, doc: Document, figure_data: Dict[str, Any]):
        """
        Add a figure to the document.

        Args:
            doc: The document
            figure_data: Figure information including number, type, title, content
        """
        fig_num = figure_data.get('number', 1)
        fig_type = figure_data.get('type', 'unknown')
        fig_title = figure_data.get('title', f'Figure {fig_num}')
        content = figure_data.get('content', '')
        content_type = figure_data.get('content_type', 'mermaid')

        print(f"Adding figure {fig_num}: {fig_title} (type: {fig_type}, content_type: {content_type})")

        # Skip if already rendered
        if fig_num in getattr(self, '_generated_figures', []):
            print(f"  Figure {fig_num} already rendered, skipping")
            return

        # Handle different content types
        image_path = None

        if content_type == 'file_path':
            # Content is already a file path (from PDF extraction or web search)
            path = Path(content)
            if path.exists():
                image_path = path
                print(f"  Using existing image: {image_path}")
            else:
                # Try to find the filepath in figure_data
                filepath = figure_data.get('filepath')
                if filepath and Path(filepath).exists():
                    image_path = Path(filepath)
                    print(f"  Using filepath from figure_data: {image_path}")
                else:
                    print(f"  File path not found: {content}")
                    return
        else:
            # Generate the image file
            print(f"  Content preview: {str(content)[:100]}...")
            try:
                output_dir = f"output/images/{self.paper_id}"
                print(f"  Creating ChartGenerator with output_dir: {output_dir}")
                chart_gen = ChartGenerator(output_dir=output_dir)

                if content_type == 'mermaid':
                    # Generate diagram from mermaid code
                    image_path = chart_gen.generate_mermaid_diagram(
                        mermaid_code=content,
                        paper_id=self.paper_id,
                        figure_num=fig_num
                    )
                elif content_type == 'chart_data':
                    # Generate chart from data
                    image_path = chart_gen.generate_matplotlib_chart(
                        data=content,
                        chart_type=fig_type,
                        title=fig_title,
                        paper_id=self.paper_id,
                        figure_num=fig_num
                    )

            except Exception as e:
                print(f"Failed to generate figure {fig_num}: {e}")
                return

        if not image_path or not image_path.exists():
            print(f"Image not available for figure {fig_num}")
            return

        # Track rendered figure
        if not hasattr(self, '_generated_figures'):
            self._generated_figures = []
        self._generated_figures.append(fig_num)

        # Add figure to document
        try:
            # Add spacing before figure
            doc.add_paragraph()

            # Add the image
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Calculate image size (max width 5.5 inches)
            run = para.add_run()
            run.add_picture(str(image_path), width=Inches(5.5))

            # Add caption
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Format caption based on language
            if self.is_chinese:
                caption_text = f"图{fig_num}  {fig_title}"
            else:
                caption_text = f"Figure {fig_num}. {fig_title}"

            caption_run = caption_para.add_run(caption_text)
            caption_run.font.size = Pt(10)
            caption_run.font.name = self.styles['normal']['font_name']

            # For Chinese papers, set East Asian font
            if self.is_chinese:
                caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), self.styles['normal']['font_name'])
                if 'font_name_latin' in self.format_config:
                    caption_run.font.name = self.format_config['font_name_latin']

            # Italic for English captions
            if not self.is_chinese:
                caption_run.italic = True

            # Add spacing after figure
            doc.add_paragraph()

        except Exception as e:
            print(f"Failed to insert figure {fig_num}: {e}")

    def convert_to_pdf(self, docx_path: Path, pdf_path: Optional[Path] = None) -> Path:
        """
        Convert DOCX to PDF.

        Args:
            docx_path: Path to the DOCX file
            pdf_path: Optional output path for PDF, defaults to same name with .pdf extension

        Returns:
            Path to the generated PDF file
        """
        if pdf_path is None:
            pdf_path = docx_path.with_suffix('.pdf')

        # Ensure output directory exists
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        # Try different conversion methods based on platform
        system = platform.system()

        if system == 'Windows':
            # Try docx2pdf on Windows (uses Microsoft Word)
            try:
                from docx2pdf import convert
                convert(str(docx_path), str(pdf_path))
                return pdf_path
            except Exception as e:
                print(f"docx2pdf failed: {e}, trying LibreOffice...")

        # Try LibreOffice (cross-platform)
        try:
            cmd = [
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(pdf_path.parent),
                str(docx_path)
            ]
            subprocess.run(cmd, check=True, capture_output=True, timeout=60)

            # LibreOffice saves as same name with .pdf extension
            expected_pdf = docx_path.with_suffix('.pdf')
            if expected_pdf.exists():
                if expected_pdf != pdf_path:
                    expected_pdf.rename(pdf_path)
                return pdf_path

        except FileNotFoundError:
            print("LibreOffice not found (soffice command)")
        except subprocess.TimeoutExpired:
            print("LibreOffice conversion timed out")
        except subprocess.CalledProcessError as e:
            print(f"LibreOffice conversion failed: {e.stderr.decode()}")

        raise RuntimeError(
            "PDF conversion failed. Please install either:\n"
            "- Microsoft Word (Windows) for docx2pdf\n"
            "- LibreOffice (cross-platform)"
        )

    def generate_pdf(
        self,
        sections: Dict[str, str],
        metadata: Dict[str, Any],
        output_path: Optional[Path] = None,
        temp_dir: Optional[Path] = None
    ) -> Path:
        """
        Generate a PDF file directly from paper sections.

        Args:
            sections: Dictionary with paper sections
            metadata: Paper metadata (title, keywords, field, etc.)
            output_path: Optional output path, defaults to output/{paper_id}.pdf
            temp_dir: Optional directory for temporary DOCX file

        Returns:
            Path to the generated PDF file
        """
        # Generate DOCX first
        if temp_dir is None:
            temp_dir = Path('output/temp')
        temp_dir.mkdir(parents=True, exist_ok=True)

        temp_docx = temp_dir / f"temp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        try:
            # Generate DOCX
            self.generate_paper(sections, metadata, temp_docx)

            # Convert to PDF
            if output_path is None:
                output_dir = Path('output')
                output_dir.mkdir(exist_ok=True)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = output_dir / f"paper_{timestamp}.pdf"

            return self.convert_to_pdf(temp_docx, output_path)

        finally:
            # Clean up temporary DOCX
            if temp_docx.exists():
                try:
                    temp_docx.unlink()
                except Exception:
                    pass

    def generate_from_paper_id_pdf(
        self,
        paper_id: str,
        papers_db: Dict[str, Dict[str, Any]]
    ) -> Optional[Path]:
        """
        Generate PDF for a paper from the database.

        Args:
            paper_id: The paper ID
            papers_db: The papers database dictionary

        Returns:
            Path to the generated PDF file or None if paper not found
        """
        if paper_id not in papers_db:
            return None

        paper = papers_db[paper_id]
        sections = paper.get('sections')
        if not sections:
            return None

        request = paper.get('request', {})
        metadata = {
            'title': paper.get('title', ''),
            'keywords': request.get('keywords', []),
            'field': request.get('field', ''),
            'citation_style': request.get('citation_style', 'apa'),
            'language': request.get('language', 'english'),
        }

        output_dir = Path('output')
        output_dir.mkdir(exist_ok=True)
        output_path = output_dir / f"{paper_id}.pdf"

        return self.generate_pdf(sections, metadata, output_path)

    def generate_from_paper_id(
        self,
        paper_id: str,
        papers_db: Dict[str, Dict[str, Any]]
    ) -> Optional[Path]:
        """
        Generate DOCX for a paper from the database.

        Args:
            paper_id: The paper ID
            papers_db: The papers database dictionary

        Returns:
            Path to the generated DOCX file or None if paper not found
        """
        if paper_id not in papers_db:
            return None

        paper = papers_db[paper_id]
        sections = paper.get('sections')
        if not sections:
            return None

        request = paper.get('request', {})
        metadata = {
            'title': paper.get('title', ''),
            'keywords': request.get('keywords', []),
            'field': request.get('field', ''),
            'citation_style': request.get('citation_style', 'apa'),
            'language': request.get('language', 'english'),
        }

        # Get figures from paper - stored in the sections dict by WriterAgent
        figures = []
        if sections and isinstance(sections, dict):
            figures = sections.get('figures', [])
            print(f"Found {len(figures)} figures in paper data")
            for fig in figures:
                print(f"  Figure {fig.get('number')}: {fig.get('title')} ({fig.get('type')})")

        output_dir = Path('output')
        output_dir.mkdir(exist_ok=True)
        output_path = output_dir / f"{paper_id}.docx"

        return self.generate_paper(sections, metadata, output_path, figures, paper_id)


def generate_docx(
    sections: Dict[str, str],
    metadata: Dict[str, Any],
    output_path: Optional[Path] = None
) -> Path:
    """
    Convenience function to generate a DOCX file.

    Args:
        sections: Dictionary with paper sections
        metadata: Paper metadata
        output_path: Optional output path

    Returns:
        Path to the generated DOCX file
    """
    generator = DocxGenerator()
    return generator.generate_paper(sections, metadata, output_path)
